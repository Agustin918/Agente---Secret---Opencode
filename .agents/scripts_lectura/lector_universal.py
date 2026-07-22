import sys, os, json, csv, xml.etree.ElementTree as ET, html.parser
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from pathlib import Path
from datetime import datetime

def sizeof_fmt(num):
    for unit in ("B", "KB", "MB", "GB"):
        if abs(num) < 1024:
            return f"{num:.1f}{unit}"
        num /= 1024
    return f"{num:.1f}TB"

def read_image(path):
    try:
        from PIL import Image
        import pytesseract
        img = Image.open(path)
        w, h = img.size
        text = pytesseract.image_to_string(img, lang="spa+eng")
        return {
            "tipo": "imagen",
            "dimensiones": f"{w}x{h}px",
            "formato": img.format,
            "modo": img.mode,
            "texto_extraido": text.strip() or "(sin texto detectable)"
        }
    except ImportError:
        return {"tipo": "imagen", "error": "Faltan librerias: pip install Pillow pytesseract"}

def read_pdf(path):
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                pages.append({"pagina": i+1, "texto": text.strip() if text else "(sin texto)"})
        return {"tipo": "pdf", "paginas": len(pages), "contenido": pages}
    except ImportError:
        return {"tipo": "pdf", "error": "Falta libreria: pip install pdfplumber"}

def read_docx(path):
    try:
        import docx
        doc = docx.Document(path)
        parrafos = [p.text for p in doc.paragraphs if p.text.strip()]
        tablas = []
        for table in doc.tables:
            filas = []
            for row in table.rows:
                filas.append([cell.text for cell in row.cells])
            tablas.append(filas)
        return {"tipo": "word", "parrafos": len(parrafos), "contenido": parrafos, "tablas": tablas}
    except ImportError:
        return {"tipo": "word", "error": "Falta libreria: pip install python-docx"}

def read_xlsx(path):
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        hojas = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            filas = []
            for row in ws.iter_rows(values_only=True):
                filas.append([str(c) if c is not None else "" for c in row])
            hojas.append({"nombre": sheet_name, "filas": len(filas), "datos": filas[:50]})
        wb.close()
        return {"tipo": "excel", "hojas": hojas}
    except ImportError:
        return {"tipo": "excel", "error": "Falta libreria: pip install openpyxl"}

def read_video(path):
    info = {"tipo": "video", "archivo": path.name}
    try:
        import subprocess
        result = subprocess.run(
            ["ffprobe", "-v", "quiet", "-print_format", "json", "-show_format", "-show_streams", str(path)],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0:
            data = json.loads(result.stdout)
            fmt = data.get("format", {})
            info["duracion"] = f"{float(fmt.get('duration', 0)):.1f}s"
            info["tamano"] = sizeof_fmt(int(fmt.get("size", 0)))
            info["codec"] = fmt.get("format_name", "")
            streams = data.get("streams", [])
            for s in streams:
                if s.get("codec_type") == "video":
                    info["resolucion"] = f"{s.get('width','?')}x{s.get('height','?')}"
                    info["fps"] = s.get("r_frame_rate", "")
                    break
        else:
            info["error"] = "ffprobe no disponible"
    except:
        info["error"] = "ffprobe no instalado"
    return info

def read_text(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
        return {"tipo": "texto", "extension": path.suffix, "lineas": content.count("\n")+1, "contenido": content[:5000]}
    except UnicodeDecodeError:
        with open(path, "r", encoding="latin-1") as f:
            content = f.read()
        return {"tipo": "texto", "extension": path.suffix, "lineas": content.count("\n")+1, "contenido": content[:5000]}

READERS = {
    ".jpg": read_image, ".jpeg": read_image, ".png": read_image,
    ".gif": read_image, ".bmp": read_image, ".tiff": read_image, ".webp": read_image,
    ".pdf": read_pdf,
    ".docx": read_docx, ".doc": read_docx,
    ".xlsx": read_xlsx, ".xls": read_xlsx,
    ".mp4": read_video, ".avi": read_video, ".mov": read_video,
    ".mkv": read_video, ".webm": read_video,
    ".txt": read_text, ".md": read_text, ".csv": read_text,
    ".json": read_text, ".xml": read_text, ".html": read_text, ".log": read_text,
}

def main():
    if len(sys.argv) < 2:
        print("USO: python lector_universal.py <ruta_del_archivo>")
        sys.exit(1)

    path = Path(sys.argv[1])
    if not path.exists():
        print(f"ERROR: archivo no encontrado: {path}")
        sys.exit(1)

    ext = path.suffix.lower()
    reader = READERS.get(ext)
    if not reader:
        print(f"EXTENSION NO SOPORTADA: {ext}")
        print(json.dumps({"tipo": "desconocido", "archivo": path.name,
                          "tamano": sizeof_fmt(path.stat().st_size),
                          "modificado": datetime.fromtimestamp(path.stat().st_mtime).isoformat()}, indent=2))
        sys.exit(0)

    info = {"archivo": path.name, "ruta": str(path), "tamano": sizeof_fmt(path.stat().st_size)}
    try:
        result = reader(path)
        info.update(result)
    except Exception as e:
        info["error"] = str(e)

    print(json.dumps(info, indent=2, ensure_ascii=False))

if __name__ == "__main__":
    main()
